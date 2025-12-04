from flask import Flask, request, render_template, send_file, session
import os
from fpdf import FPDF
from docx import Document
import docx
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from google import genai
import google.generativeai as genai2
import zipfile
import json
import pandas as pd
from pydantic import BaseModel
import matplotlib.pyplot as plt
from matplotlib.backends.backend_pdf import PdfPages
import subprocess

# Define the timeout handler
def timeout_handler(signum, frame):
    raise TimeoutError("Function execution timed out.")
class tableschema(BaseModel):
  table_names: str
  table_datas: str

#--------------------------------------------------------------------------------------------------
client = genai.Client(api_key="add your gemini api here")
model_id = "gemini-2.0-flash"
genai2.configure(api_key="and here ")


app = Flask(__name__)
app.config['SEND_FILE_MAX_AGE_DEFAULT'] = 0  # Disable cache
app.config['MAX_CONTENT_LENGTH'] = 100 * 1024 * 1024
app.secret_key = '''i don't remember why i added that but add secret key here '''

# ----------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------
DEFAULT_LANGUAGE = "Extract Text (up to 10 images at once)"
# ----------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------


@app.route('/file_type', methods=['GET', 'POST'])
def Ty():
    if 'language' in request.form:
        language = request.form.get('language')
        session['language'] = language
    print("language is:",session['language'])
    if request.method == 'POST':
        session['file_type'] = request.form.get('filetype')
    return render_template("index.html", selected_language=session['language'],selected_option=session['file_type'])

# ----------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------

# ----------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------
@app.route("/", methods=['GET', 'POST'])
def main():
    session.setdefault('language', DEFAULT_LANGUAGE)
    session['file_type'] = 'PDF'
    session['language'] = DEFAULT_LANGUAGE
    print("language is:",session['language'])
    return render_template("index.html", selected_language=session['language'],selected_option=session['file_type'])

# ----------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------

@app.route("/preview_text", methods=['GET', 'POST'])

def preview():
    try:
        file_type = session['file_type']
        print("preview_text",file_type)
        if 'language' in request.form and request.form.get('language') != "":
            language = request.form.get('language')
            session['language'] = language
        else:
            session['language'] = DEFAULT_LANGUAGE
        print("language is:",session['language'])
        if session['language']=="Image description (up to 10 images at once)":
            def processing():

                import google.generativeai as genai

                genai.configure(api_key="AIzaSyC6dIrUxmGDoU7fwhfLwfpeBwUUuQqnmvg")

                def upload_to_gemini(path, mime_type=None):
                    """Uploads the given file to Gemini.

                    See https://ai.google.dev/gemini-api/docs/prompting_with_media
                    """
                    file = genai.upload_file(path, mime_type=mime_type)
                    print(f"Uploaded file '{file.display_name}' as: {file.uri}")
                    return file

                content = []
                uploaded_files = request.files.getlist('my_images')


                for file_path in uploaded_files:
                    file_path2 = file_path.filename
                    session['image_names'] = [file_path2]
                    file_path.save(file_path2)
                    file_response = upload_to_gemini(file_path2, mime_type="image/jpeg")
                    content.append(file_response)
                content.append("""Follow this instructions:
                               1-can you provide detailed description of the image(s) without any extra text.
                               2-if the image in language other than english descripe in the other language.
                               3-if there are more than one picture number the pictures and descripe each one of them.
                               4-Extract the text exactly as it appears in the image. Do not add any introductions, explanations, comments, or extra formatting. Always return only the extracted text with no additional words. This instruction must be followed strictly""")
                # Create the model
                generation_config = {
                "temperature": 1,
                "top_p": 0.95,
                "top_k": 40,
                "max_output_tokens": 8192,
                "response_mime_type": "text/plain",
                }

                model = genai.GenerativeModel(
                model_name="gemini-2.0-flash",
                generation_config=generation_config,
                )

                chat_session = model.start_chat(
                )

                response = chat_session.send_message(content)
                session['texts'] = response.text
                return response.text


            result = processing()
            print('done')
    # ----------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------

        elif session['language']== "Extract Text (up to 10 images at once)":
            def processing():

                import google.generativeai as genai

                genai.configure(api_key="AIzaSyC6dIrUxmGDoU7fwhfLwfpeBwUUuQqnmvg")

                def upload_to_gemini(path, mime_type=None):
                    """Uploads the given file to Gemini.

                    See https://ai.google.dev/gemini-api/docs/prompting_with_media
                    """
                    file = genai.upload_file(path, mime_type=mime_type)
                    print(f"Uploaded file '{file.display_name}' as: {file.uri}")
                    return file

                content = []
                uploaded_files = request.files.getlist('my_images')


                for file_path in uploaded_files:
                    file_path2 = file_path.filename
                    session['image_names'] = [file_path2]
                    file_path.save(file_path2)
                    file_response = upload_to_gemini(file_path2, mime_type="image/jpeg")
                    content.append(file_response)
                content.append("""
                               extract the text in the images

                               """)
                # Create the model
                generation_config = {
                "temperature": 1,
                "top_p": 0.95,
                "top_k": 40,
                "max_output_tokens": 8192,
                "response_mime_type": "text/plain",
                }

                model = genai.GenerativeModel(
                model_name="gemini-2.0-flash-thinking-exp-01-21",
                system_instruction="""
                               1-please provide the text in the images
                               2-please make sure that each line in the image is in seperate line of the text without any extra text at all else say There is no text in the image.
                               3-number the pictures please.
                               4-please Extract the text exactly as it appears in the image. Do not add any introductions, explanations, comments, or extra formatting. Always return only the extracted text with no additional words. This instruction must be followed strictly as you are ocr model

                               """,
                generation_config=generation_config,
                )

                chat_session = model.start_chat(
                )

                response = chat_session.send_message(content)
                session['texts'] = response.text
                return response.text

            result = processing()
            print('done')
    # ----------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------

        elif session['language']=="Extract Tables (up to 10 images at once)":
            def processing():

                import google.generativeai as genai
                genai.configure(api_key="AIzaSyC6dIrUxmGDoU7fwhfLwfpeBwUUuQqnmvg")
                def upload_to_gemini(path, mime_type=None):
                    """Uploads the given file to Gemini.

                    See https://ai.google.dev/gemini-api/docs/prompting_with_media
                    """
                    file = genai.upload_file(path, mime_type=mime_type)
                    print(f"Uploaded file '{file.display_name}' as: {file.uri}")
                    return file

                content = []
                uploaded_files = request.files.getlist('my_images')
                print(uploaded_files)

                for file_path in uploaded_files:
                    file_path2 = file_path.filename
                    session['image_names'] = [file_path2]
                    file_path.save(file_path2)
                    file_response = upload_to_gemini(file_path2, mime_type="image/jpeg")
                    content.append(file_response)
                content.append("provide me with all tables in the images")
                generation_config = {
                    "temperature": 1,
                    "top_p": 0.95,
                    "top_k": 40,
                    "max_output_tokens": 8192,
                    "response_mime_type": "application/json",
                    }

                model = genai.GenerativeModel(
                system_instruction="""
                                Provide JSON representations of all tables present in the input
                                Note: Ensure that all columns in the same table have the same number of rows (data entries).
                                Ensure that all columns in the same table have the same number of rows (data entries).
                                Ensure that all columns in the same table have the same number of rows (data entries).
                                Ensure that all columns in the same table have the same number of rows (data entries).
                                Ensure that all columns in the same table have the same number of rows (data entries).
                                Ensure that all columns in the same table have the same number of rows (data entries).
                                All arrays must be of the same length
                                All arrays must be of the same length
                                All arrays must be of the same length
                                All arrays must be of the same length
                                All arrays must be of the same length
                                  Return accurate results and if there is no column name no problem add column name as " ". Use the following JSON schema:

                                    {
                                    "table_name": "name_of_the_first_table",
                                    "table_data": {
                                    "column1": ["data1", "data2", "data3"],
                                    "column2": ["data1", "data2", "data3"]
                                                }
                                    },
                                    {
                                    "table_name": "name_of_the_second_table",
                                    "table_data": {
                                    "column1": ["data1", "data2"],
                                    "column2": ["data1", "data2"]
                                                }
                                    }
                               """,
                model_name="gemini-2.0-flash",
                generation_config=generation_config,
                )

                chat_session = model.start_chat(
                )

                response = chat_session.send_message(content)
                session['texts'] = [response.text]
                print("text is",response.text)
                return response.text


            result = processing()
            print('done')
    # ----------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------

        elif session['language']=="Extract Math (up to 10 images at once)":
            if file_type == 'DOC':
                def processing():

                    import google.generativeai as genai
                    genai.configure(api_key="AIzaSyC6dIrUxmGDoU7fwhfLwfpeBwUUuQqnmvg")
                    def upload_to_gemini(path, mime_type=None):
                        """Uploads the given file to Gemini.

                        See https://ai.google.dev/gemini-api/docs/prompting_with_media
                        """
                        file = genai.upload_file(path, mime_type=mime_type)
                        print(f"Uploaded file '{file.display_name}' as: {file.uri}")
                        return file

                    content = []
                    uploaded_files = request.files.getlist('my_images')
                    print(uploaded_files)

                    for file_path in uploaded_files:
                        file_path2 = file_path.filename
                        session['image_names'] = [file_path2]
                        file_path.save(file_path2)
                        file_response = upload_to_gemini(file_path2, mime_type="image/jpeg")
                        content.append(file_response)
                    content.append("provide me with all tables in the images")
                    generation_config = {
                        "temperature": 1,
                        "top_p": 0.95,
                        "top_k": 40,
                        "max_output_tokens": 8192
                        }

                    model = genai.GenerativeModel(
                    system_instruction="""solve this equation(s) by explaining through latex only with no text at all in the middle like since or therefor etc.. just numbers and sympols, and begin with $$ and end with $$, never use one dollar sign ($) and don't put all equations at one line each equation=something in a line by itself  (if there is no equations say There are no equations),If there are more than one problem number them as problem and it's number then the problem then solution and it's number then the problem's solution, there must be only two $$ like $$ problem1 ..... solution1.... problem2 ..... solution2.... $$,The image(s) will be displayed on webpage so it can't be displayed if there is anything outside $$ ....$$(never put things like \\documentclass{article} \\usepackage{amsmath} \\usepackage{amsfonts} \\usepackage{amssymb} \\begin{document} \\end{document})""",
                    model_name="gemini-2.0-flash-thinking-exp-01-21",
                    generation_config=generation_config,
                    )

                    chat_session = model.start_chat(
                    )

                    response = chat_session.send_message(content)
                    session['texts'] = [response.text]
                    print("text is",response.text)

                    solved_latex = response.text.replace("```","")
                    solved_latex = solved_latex.replace("latex","").replace("\\\\","\\")
                    return solved_latex

                result = processing()
                session['texts']=result
                print(result)
            else:
                def processing():

                    import google.generativeai as genai
                    genai.configure(api_key="AIzaSyC6dIrUxmGDoU7fwhfLwfpeBwUUuQqnmvg")
                    def upload_to_gemini(path, mime_type=None):
                        """Uploads the given file to Gemini.

                        See https://ai.google.dev/gemini-api/docs/prompting_with_media
                        """
                        file = genai.upload_file(path, mime_type=mime_type)
                        print(f"Uploaded file '{file.display_name}' as: {file.uri}")
                        return file

                    content = []
                    uploaded_files = request.files.getlist('my_images')
                    print(uploaded_files)

                    for file_path in uploaded_files:
                        file_path2 = file_path.filename
                        session['image_names'] = [file_path2]
                        file_path.save(file_path2)
                        file_response = upload_to_gemini(file_path2, mime_type="image/jpeg")
                        content.append(file_response)
                    content.append("provide me with all tables in the images")
                    generation_config = {
                        "temperature": 1,
                        "top_p": 0.95,
                        "top_k": 40,
                        "max_output_tokens": 8192
                        }

                    model = genai.GenerativeModel(
                    system_instruction="""extract math equations through latex only with no text at all in the middle like since or therefor etc.. just numbers and sympols, and begin with $$ and end with $$, never use one dollar sign ($) and don't put all equations at one line each equation=something in a line by itself (if there is no equations say There are no equations),If there are more than one problem number them as problem and it's number then the problem, there must be only two $$ like $$ problem1 ..... .... problem2 ..... .... $$,The image(s) will be displayed on webpage so it can't be displayed if there is anything outside $$ ....$$(never put things like \\documentclass{article} \\usepackage{amsmath} \\usepackage{amsfonts} \\usepackage{amssymb} \\begin{document} \\end{document})""",
                    model_name="gemini-2.0-flash-thinking-exp-01-21",
                    generation_config=generation_config,
                    )

                    chat_session = model.start_chat(
                    )

                    response = chat_session.send_message(content)
                    session['texts'] = [response.text]
                    print("text is",response.text)

                    solved_latex = response.text.replace("```","")
                    solved_latex = solved_latex.replace("latex","").replace("\\\\","\\")
                    return solved_latex

                result = processing()
                session['texts']=result
                print(result)
    # ----------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------

        elif session['language']=="Extract Information (up to 10 images at once)":


            def processing():
                user_input = request.form.get('input')
                print("vsdvdsvdsvsvsvsd",user_input)
                import google.generativeai as genai
                genai.configure(api_key="AIzaSyC6dIrUxmGDoU7fwhfLwfpeBwUUuQqnmvg")
                def upload_to_gemini(path, mime_type=None):
                    """Uploads the given file to Gemini.

                    See https://ai.google.dev/gemini-api/docs/prompting_with_media
                    """
                    file = genai.upload_file(path, mime_type=mime_type)
                    print(f"Uploaded file '{file.display_name}' as: {file.uri}")
                    return file

                content = []
                uploaded_files = request.files.getlist('my_images')
                print(uploaded_files)

                for file_path in uploaded_files:
                    file_path2 = file_path.filename
                    session['image_names'] = [file_path2]
                    file_path.save(file_path2)
                    file_response = upload_to_gemini(file_path2, mime_type="image/jpeg")
                    content.append(file_response)
                content.append(f"what the user want is : {user_input}")
                generation_config = {
                    "temperature": 1,
                    "top_p": 0.95,
                    "top_k": 40,
                    "max_output_tokens": 8192
                    }

                model = genai.GenerativeModel(
                system_instruction=""" process every image and extract the next information from the images in details and answer only in the language the user speakes after the next simicolumn with no extra text at all the user doesn't know you are a chatbot""",
                model_name="gemini-2.0-flash-thinking-exp-01-21",
                generation_config=generation_config,
                )

                chat_session = model.start_chat(
                )

                response = chat_session.send_message(content)
                session['texts'] = [response.text]
                print("text is",response.text)

                return response.text

            result = processing()
            session['texts']=result
            print(result)
        return result
    # ----------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------
    except:
        result = "Error processing image(s)"
        return result
# ----------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------
@app.route("/Download_button", methods=['POST'])
def get_output1():
    try:

        if 'language' in request.form:
            language = request.form.get('language')
            session['language'] = language
        print("language is:",session['language'])
        if request.method == 'POST':

            if request.form.get('file_type') is not None:
                session['file_type'] = request.form.get('file_type')

            if session['file_type'] == 'PDF':
                print("t3eeeeeeeeeeeeeeeeeeeent  ",session['language'])

                if session['language'] == "Extract Tables (up to 10 images at once)":

                    json_string = session['texts']
                    print(json_string)
                    if type(json_string) is list:
                        json_string = json_string[0]
                        print("a list")

                    json_data = json.loads(json_string)
                    print(json_data)
                    # Create a PDF

                    file_name2 = session['image_names']
                    file_name2 = f"{file_name2}.pdf"
                    file_path = os.path.abspath(file_name2)
                    print(range(len(json_data)))
                    with PdfPages(file_path) as pp:

                        for i in range(len(json_data)):
                            try:
                                table_name = json_data[i]['table_name'].replace('/','_').replace("\\","_" )
                                table_data = json_data[i]['table_data']
                                max_length = max(len(values) for values in table_data.values())

                                filled_data = {key: values + ["NULL"] * (max_length - len(values)) for key, values in table_data.items()}
                                df = pd.DataFrame(filled_data)


                                fig, ax =plt.subplots(figsize=(12,4))
                                ax.text(0.5, 1.15, table_name, fontsize='x-large', color="black", ha="center", transform=ax.transAxes)
                                ax.axis('tight')
                                ax.axis('off')
                                the_table = ax.table(cellText=df.values,colLabels=df.columns,loc='center')
                                pp.savefig( bbox_inches='tight')

                            except Exception as e:
                                print("you failed me bot")
                                print(str(e))

                        pp.close()
                    return send_file(file_path, as_attachment=True, download_name=file_name2)
                elif session['language'] == "Extract Math (up to 10 images at once)":
                    def tex_to_pdf(tex_file_path, output_dir=None):
                        output_dir = output_dir or os.path.dirname(tex_file_path)

                        # Run pdflatex to compile the .tex file
                        try:
                            result = subprocess.run(
                                ["pdflatex", "-output-directory", output_dir, tex_file_path],
                                check=True,
                                text=True,
                                stdout=subprocess.PIPE,
                                stderr=subprocess.PIPE,
                            )
                            print("PDF generated successfully!")
                            print(result.stdout)
                        except subprocess.CalledProcessError as e:
                            print(f"Error during PDF generation: {e} ")

                    math_latex = session['texts']

                    math_latex = r"""\pdfminorversion=4
                                    \documentclass[]{article}
                                    \usepackage[utf8]{inputenc}
                                    \usepackage{amssymb,latexsym,amsmath}
                                    \usepackage{graphicx}
                                    \usepackage[colorlinks=true, allcolors=blue]{hyperref}
                                    \begin{document}
                                    \Large """ +math_latex + r"""\end {document}"""

                    file_name = session['image_names']
                    file_name2= f"{file_name}.pdf"
                    file_name3 = f"{file_name}.tex"
                    file_path =os.path.abspath(file_name2)
                    file_path2 =os.path.abspath(file_name3)
                    if os.path.exists(file_path):
                        os.renames(file_path,f"{file_path}_old")
                        os.remove(f"{file_path}_old")
                    if os.path.exists(file_path2):
                        os.renames(file_path2,f"{file_path2}_old")
                        os.remove(f"{file_path2}_old")

                    with open(file_path2, "wb") as f:
                        f.write(math_latex.encode('utf-8'))
                    tex_to_pdf(file_path2 ,output_dir=os.path.dirname(file_path2))

                    return send_file(file_path, as_attachment=True, download_name=file_name2)

                else:
                    pdf = FPDF()
                    pdf.set_auto_page_break(auto=True, margin=15)
                    font_path = 'arial-unicode-ms.ttf'
                    pdf.add_font('arial-unicode-ms', '', font_path, uni=True)
                    pdf.set_font('arial-unicode-ms', size=12)

                    for text1 in [session['texts']]:
                        pdf.add_page()
                        pdf.multi_cell(0, 8, txt=text1)
                    file_name = session['image_names'][0]
                    file_name2= f"{file_name}.pdf"
                    file_path = os.path.abspath(file_name2)
                    pdf.output(file_path)

                    return send_file(file_path, as_attachment=True, download_name=file_name2)
    ##########################################################################################################################################################################
            elif session['file_type'] == 'TXT':

                # Define the file path for the zip file
                if session['language'] == "Extract Tables (up to 10 images at once)":

                    file_name2 = session['image_names']
                    file_name2 = f"{file_name2}.zip"
                    file_path = os.path.abspath(file_name2)
                    json_string = session['texts']
                    print("json type is ",type(json_string))
                    if type(json_string) is list:
                        json_string = json_string[0]
                        print("a list")
                    json_data = json.loads(json_string)
                    with zipfile.ZipFile(file_path, 'w', zipfile.ZIP_DEFLATED) as zip_file:
                            for i in range(len(json_data)):
                                try:
                                    table_name = json_data[i]['table_name'].replace('/','_').replace("\\","_" )
                                    table_data = json_data[i]['table_data']
                                    df = pd.DataFrame(table_data)
                                    file_path2 = f"{table_name}.xlsx"
                                    file_path3 = os.path.abspath(file_path2)
                                    df.to_excel(file_path2)
                                    if os.path.exists(file_path2):
                                        zip_file.write(os.path.basename(file_path2), arcname=file_path2)
                                except Exception as e:
                                    print("you failed me bot")
                                    print(str(e))

                    if len(json_data) > 1:
                        return send_file(file_path, as_attachment=True, download_name=file_name2)
                    else:
                        return send_file(file_path3, as_attachment=True, download_name=file_path2)
    #############################################################################################################################################################################################
                elif session['language'] == "Extract Math (up to 10 images at once)":

                    math_latex = session['texts']
                    math_latex = r"""\pdfminorversion=4
                                    \documentclass[]{article}
                                    \usepackage[utf8]{inputenc}
                                    \usepackage{amssymb,latexsym,amsmath}
                                    \usepackage{graphicx}
                                    \usepackage[colorlinks=true, allcolors=blue]{hyperref}
                                    \begin{document}
                                    \Large """ +math_latex + r"""\end {document}"""

                    file_name2 = session['image_names']
                    file_name2 = f"{file_name2}.tex"
                    file_path = os.path.abspath(file_name2)
                    with open(f"{file_name2}.tex", "wb") as f:
                        f.write(math_latex.encode('utf-8'))


                    return send_file(file_path, as_attachment=True, download_name=file_name2)
    #############################################################################################################################################################################################
                else:
                    file_name = session['image_names']
                    file_name3 = f"{file_name[0]}.zip"
                    file_path = os.path.abspath(file_name3)
                    num_images = len(session['image_names'])

                    textss = [session["texts"]]

                    print("number of images :",num_images)
                    with zipfile.ZipFile(file_path, 'w', zipfile.ZIP_DEFLATED) as zip_file:
                        for text1, file_name in zip(textss, session['image_names']):
                            # Generate file name and file path for the text file
                            file_name2 = f"{file_name}.txt"
                            file_path2 = os.path.abspath(file_name2)
                            print('cascasc',file_name)
                            # Write the text content to a file

                            with open(file_path2, "w",encoding="utf-8") as file:
                                file.write(text1)


                            # Add the file to the zip archive
                            if os.path.exists(file_path2):
                                zip_file.write(file_path2, arcname=os.path.basename(file_path2))

  ##############################################################################################################################################

                    if len(session['image_names']) > 1:
                        return send_file(file_path, as_attachment=True, download_name=file_name3)
                    else:
                        return send_file(file_path2, as_attachment=True, download_name=file_name2)
                    # Assign the zip file path to the file_path variable


    #############################################################################################################################################################################################
            elif session['file_type'] == 'DOC':
                if session['language'] == "Extract Math (up to 10 images at once)":
                    def tex_to_pdf(tex_file_path, output_dir=None):
                        output_dir = output_dir or os.path.dirname(tex_file_path)

                        # Run pdflatex to compile the .tex file
                        try:
                            result = subprocess.run(
                                ["pdflatex", "-output-directory", output_dir, tex_file_path],
                                check=True,
                                text=True,
                                stdout=subprocess.PIPE,
                                stderr=subprocess.PIPE,
                            )
                            print("PDF generated successfully!")
                        except subprocess.CalledProcessError as e:
                            print(f"Error during PDF generation:{e}")

                    math_latex = session['texts']
                    math_latex = r"""\pdfminorversion=4
                                    \documentclass[]{article}
                                    \usepackage[utf8]{inputenc}
                                    \usepackage{amssymb,latexsym,amsmath}
                                    \usepackage{graphicx}
                                    \usepackage[colorlinks=true, allcolors=blue]{hyperref}
                                    \begin{document}
                                    \Large """ +math_latex + r"""\end {document}"""

                    file_name = session['image_names']
                    file_name2= f"{file_name}.pdf"
                    file_name3 = f"{file_name}.tex"
                    file_path =os.path.abspath(file_name2)
                    file_path2 =os.path.abspath(file_name3)
                    if os.path.isfile(file_path):
                        os.renames(file_path,f"{file_path}_old")
                        os.remove(f"{file_path}_old")
                    if os.path.isfile(file_path2):
                        os.renames(file_path2,f"{file_path2}_old")
                        os.remove(f"{file_path2}_old")
                    with open(file_path2, "wb") as f:
                        f.write(math_latex.encode('utf-8'))
                    tex_to_pdf(file_path2 ,output_dir=os.path.dirname(file_path2))

                    return send_file(file_path, as_attachment=True, download_name=file_name2)

                elif session['language'] == "Extract Tables (up to 10 images at once)":
                    doc = Document()

                    doc.styles['Normal'].font.name = 'Arial Unicode MS'
                    doc.styles['Normal'].font.size = Pt(12)
                    file_name2 = session['image_names']
                    file_name2 = f"{file_name2}.docx"
                    file_path = os.path.abspath(f'{file_name2}.docx')
                    json_string = session['texts']
                    if type(json_string) is list:
                        json_string = json_string[0]
                        print("a list")
                    json_data = json.loads(json_string)
                    for i in range(len(json_data)):
                        try:
                            table_name = json_data[i]['table_name'].replace('/','_').replace("\\","_" )
                            table_data = json_data[i]['table_data']
                            df = pd.DataFrame(table_data)
                            title = doc.add_paragraph()
                            title.alignment = 1  # Center align the title
                            run = title.add_run(table_name)
                            run.bold = True
                            run.font.size = Pt(14)

                            t = doc.add_table(df.shape[0]+1, df.shape[1])

                            # add the header rows.
                            for j in range(df.shape[-1]):
                                t.cell(0,j).text = df.columns[j]

                            # add the rest of the data frame
                            for i in range(df.shape[0]):
                                for j in range(df.shape[-1]):
                                    t.cell(i+1,j).text = str(df.values[i,j])
                            for row in t.rows:
                                for cell in row.cells:
                                    # Get the XML of the cell
                                    cell._element.get_or_add_tcPr().append(
                                        docx.oxml.parse_xml(
                                            r'<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                                            r'<w:top w:val="single" w:sz="4" w:color="000000"/>'
                                            r'<w:left w:val="single" w:sz="4" w:color="000000"/>'
                                            r'<w:bottom w:val="single" w:sz="4" w:color="000000"/>'
                                            r'<w:right w:val="single" w:sz="4" w:color="000000"/>'
                                            r'</w:tcBorders>'
                                        )
                                    )
                        except Exception as e:
                            print("you failed me bot")
                            print(str(e))

                        # save the doc
                    doc.save(file_path)
                    return send_file(file_path, as_attachment=True, download_name=file_name2)
                else:
                    doc = Document()
                    for text1 in [session['texts']]:

                        paragraph = doc.add_paragraph()
                        run = paragraph.add_run(text1)
                        font = run.font
                        font.name = 'Arial Unicode MS'

                        # Optionally set other paragraph properties
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # Adjust  alignment as needed

                        # Add page break (optional)
                        doc.add_page_break()
                    file_name2 = session['image_names'][0]
                    print("vdvsvdfilenamesession names imagese",session['image_names'])
                    print("vdvsvdfilenamefilename",file_name2)
                    file_name2 = f"{file_name2}.docx"
                    file_path = os.path.abspath(file_name2)
                    doc.save(file_path)
                    return send_file(file_path, as_attachment=True, download_name=file_name2)





    except Exception as e:
        error ="There was unexcepected error please try again later"
        file_path = os.path.abspath("ERROR.txt")
        file_name2 = "ERROR.txt"
        with open(file_path, mode="w", encoding="utf-8" ) as f:
            f.write(error)
        print(e)

        return send_file(file_path, as_attachment=True, download_name=file_name2)



    return render_template("index.html",selected_language=session['language'],selected_option=session['file_type'])
# ----------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------
@app.route('/privacy_policy')
def privacy_policy():
    return render_template('policy.html')

@app.route('/about_us')
def about_us():
    return render_template('about.html')

@app.route('/terms_of_use')
def terms_of_use():
    return render_template('terms.html')

@app.route('/contact_us')
def contact_us():
    return render_template('contact.html')

@app.route('/robots.txt')
def robots_txt():
    return send_file('robots.txt')

@app.route('/sitemap.xml')
def site_map():
    return send_file('sitemap.xml')


if __name__ == "__main__":
    app.run(host='0.0.0.0', port=1000)